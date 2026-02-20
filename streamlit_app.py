import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pdfplumber
import re
import io
import tempfile
from docx import Document as DocxReader

st.set_page_config(page_title="Resume Formatter", layout="centered")
st.title("Resume Formatter Application")

uploaded_file = st.file_uploader("Upload Resume (Any Format)", type=None)

# ---------------------------------------------------
# 1️⃣ CONVERT ANY FILE TO CLEAN TXT
# ---------------------------------------------------
def convert_to_text(file):
    name = file.name.lower()

    if name.endswith(".pdf"):
        text = ""
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                content = page.extract_text()
                if content:
                    text += content + "\n"
        return text

    elif name.endswith(".docx"):
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file.read())
            tmp_path = tmp.name

        doc = DocxReader(tmp_path)
        return "\n".join([para.text for para in doc.paragraphs])

    elif name.endswith(".txt"):
        return file.read().decode("utf-8")

    else:
        return ""


# ---------------------------------------------------
# 2️⃣ CLEAN TEXT
# ---------------------------------------------------
def clean_text(text):
    text = re.sub(r"http\S+", "", text)  # remove URLs
    text = re.sub(r"[•●▪◦*-]", "", text)  # remove bullets/symbols
    text = re.sub(r"\t", " ", text)
    text = re.sub(r"\n+", "\n", text)
    return text.strip()


# ---------------------------------------------------
# 3️⃣ DETECT SECTIONS
# ---------------------------------------------------
def detect_sections(text):
    sections = {
        "summary": [],
        "skills": [],
        "education": [],
        "certification": [],
        "training": [],
        "experience": []
    }

    current = None
    lines = text.split("\n")

    for line in lines:
        lower = line.lower().strip()

        if "summary" in lower or "profile" in lower:
            current = "summary"
            continue
        elif "technical skills" in lower or lower == "skills":
            current = "skills"
            continue
        elif "education" in lower:
            current = "education"
            continue
        elif "certification" in lower:
            current = "certification"
            continue
        elif "training" in lower:
            current = "training"
            continue
        elif "experience" in lower:
            current = "experience"
            continue

        if current and line.strip():
            sections[current].append(line.strip())

    return sections


# ---------------------------------------------------
# 4️⃣ BUILD FINAL DOCX STRICTLY
# ---------------------------------------------------
def build_document(candidate_name, sections):

    doc = Document()

    # Global font rule
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(10)

    # ---------------- NAME ----------------
    name_para = doc.add_paragraph()
    run = name_para.add_run(candidate_name.title())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Times New Roman"
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")  # one line space

    # ---------------- SUMMARY ----------------
    if sections["summary"]:
        h = doc.add_paragraph()
        run = h.add_run("Summary")
        run.bold = True
        run.font.size = Pt(10)

        for line in sections["summary"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.add_run("• " + line.strip())

        doc.add_paragraph("")

    # ---------------- TECHNICAL SKILLS ----------------
    if sections["skills"]:
        h = doc.add_paragraph()
        run = h.add_run("Technical Skills")
        run.bold = True
        run.font.size = Pt(10)

        for line in sections["skills"]:
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

        doc.add_paragraph("")

    # ---------------- EDUCATION / CERTIFICATION / TRAINING ----------------
    if sections["education"] or sections["certification"] or sections["training"]:
        h = doc.add_paragraph()
        run = h.add_run("Education, Certification & Training")
        run.bold = True
        run.font.size = Pt(10)

        for key in ["education", "certification", "training"]:
            for line in sections[key]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.add_run("  • " + line.strip())

        doc.add_paragraph("")

    # ---------------- PROFESSIONAL EXPERIENCE ----------------
    if sections["experience"]:
        h = doc.add_paragraph()
        run = h.add_run("Professional Experience")
        run.bold = True
        run.font.size = Pt(10)

        for line in sections["experience"]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.add_run("  • " + line.strip())

    return doc


# ---------------------------------------------------
# MAIN EXECUTION
# ---------------------------------------------------
if uploaded_file:

    raw_text = convert_to_text(uploaded_file)

    if not raw_text:
        st.error("Unsupported file format.")
    else:
        cleaned = clean_text(raw_text)
        sections = detect_sections(cleaned)

        lines = cleaned.split("\n")
        candidate_name = lines[0].strip()

        final_doc = build_document(candidate_name, sections)

        buffer = io.BytesIO()
        final_doc.save(buffer)
        buffer.seek(0)

        file_name = f"{candidate_name.title()}.docx"

        st.success("Formatted Resume Ready")

        st.download_button(
            "Download Formatted Resume",
            buffer,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
